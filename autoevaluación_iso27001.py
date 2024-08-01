import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
import numpy as np
import os

# Definir las descripciones de las rúbricas específicas para cada pregunta
rubricas = {
    'Gestión de Acceso': {
        '¿Existen políticas y procedimientos documentados para la gestión de accesos?': {
            1: 'No se tienen políticas ni procedimientos documentados. Esto implica una falta de controles formales que pongan en riesgo la confidencialidad, integridad y disponibilidad de la información.',
            2: 'Existen políticas y procedimientos, pero no están completamente documentados o actualizados, lo que puede llevar a inconsistencias en su aplicación.',
            3: 'Políticas y procedimientos documentados y regularmente revisados. Se abordan la mayoría de los requisitos de la norma ISO 27001.',
            4: 'Cumplen con todos los requisitos establecidos por ISO 27001, incluyendo la documentación completa y actualizada, además de revisiones periódicas.',
            5: 'Implementación avanzada que supera los requisitos estándar, con controles adicionales y medidas proactivas para mejorar continuamente la gestión de accesos.'
        },
        '¿Se implementan controles de autenticación fuertes para acceder a sistemas críticos?': {
            1: 'No se implementan controles de autenticación, lo que expone a los sistemas críticos a accesos no autorizados.',
            2: 'Se implementan controles de autenticación de manera limitada o inconsistente, lo que puede resultar en brechas de seguridad.',
            3: 'Controles de autenticación fuertes implementados de manera regular, cubriendo la mayoría de los sistemas críticos.',
            4: 'Cumple totalmente con los requisitos de autenticación de ISO 27001, asegurando un acceso seguro a todos los sistemas críticos.',
            5: 'Implementación avanzada de controles de autenticación que supera los requisitos estándar, incluyendo autenticación multifactor y revisiones periódicas de acceso.'
        }
    },
    'Seguridad Física y Ambiental': {
        '¿Existen medidas de seguridad física para proteger los equipos críticos del departamento de sistemas?': {
            1: 'No hay medidas de seguridad física implementadas, dejando los equipos críticos vulnerables a accesos no autorizados y daños físicos.',
            2: 'Medidas de seguridad física parciales o insuficientes, lo que puede resultar en vulnerabilidades.',
            3: 'Medidas de seguridad física implementadas regularmente, protegiendo la mayoría de los equipos críticos.',
            4: 'Cumple totalmente con los requisitos de seguridad física de ISO 27001, asegurando la protección completa de los equipos críticos.',
            5: 'Implementación avanzada que supera los requisitos estándar, con controles físicos adicionales y monitoreo continuo.'
        },
        '¿Se realizan controles ambientales para proteger la infraestructura tecnológica (temperatura, humedad, etc.)?': {
            1: 'No se realizan controles ambientales, lo que puede llevar a fallos en la infraestructura tecnológica.',
            2: 'Controles ambientales realizados de manera irregular o insuficiente, exponiendo la infraestructura a riesgos ambientales.',
            3: 'Controles ambientales implementados regularmente, cubriendo la mayoría de los riesgos ambientales.',
            4: 'Cumple totalmente con los requisitos de controles ambientales de ISO 27001, asegurando un ambiente controlado y seguro.',
            5: 'Implementación avanzada que supera los requisitos estándar, incluyendo sistemas de monitoreo continuo y medidas preventivas adicionales.'
        }
    },
    'Gestión de Comunicaciones y Operaciones': {
        '¿Se utilizan procedimientos seguros para la transmisión de datos sensibles dentro y fuera de la organización?': {
            1: 'No se utilizan procedimientos seguros para la transmisión de datos, exponiendo información sensible a interceptaciones.',
            2: 'Procedimientos seguros utilizados de manera parcial o inconsistente, lo que puede resultar en vulnerabilidades.',
            3: 'Procedimientos seguros utilizados regularmente, protegiendo la mayoría de las transmisiones de datos sensibles.',
            4: 'Cumple totalmente con los requisitos de seguridad de transmisión de datos de ISO 27001, asegurando la protección de toda la información sensible transmitida.',
            5: 'Implementación avanzada que supera los requisitos estándar, con medidas adicionales como cifrado avanzado y monitoreo continuo de transmisiones.'
        },
        '¿Se realizan pruebas periódicas de vulnerabilidades y evaluaciones de riesgos en la infraestructura de redes?': {
            1: 'No se realizan pruebas de vulnerabilidades ni evaluaciones de riesgos, dejando la infraestructura de redes expuesta.',
            2: 'Pruebas de vulnerabilidades realizadas de manera limitada o irregular, lo que puede resultar en riesgos no identificados.',
            3: 'Pruebas de vulnerabilidades y evaluaciones de riesgos realizadas regularmente, cubriendo la mayoría de la infraestructura de redes.',
            4: 'Cumple totalmente con los requisitos de pruebas y evaluaciones de ISO 27001, asegurando una infraestructura de redes segura.',
            5: 'Implementación avanzada que supera los requisitos estándar, con pruebas continuas y análisis de riesgos en profundidad.'
        }
    },
    'Control de Acceso a la Información': {
        '¿Se implementan controles para limitar el acceso a la información confidencial y crítica dentro del departamento de sistemas?': {
            1: 'No se implementan controles de acceso a la información, exponiendo datos confidenciales y críticos a accesos no autorizados.',
            2: 'Controles de acceso implementados de manera limitada o inconsistente, lo que puede resultar en brechas de seguridad.',
            3: 'Controles de acceso implementados regularmente, protegiendo la mayoría de la información confidencial y crítica.',
            4: 'Cumple totalmente con los requisitos de control de acceso de ISO 27001, asegurando la protección completa de la información confidencial y crítica.',
            5: 'Implementación avanzada que supera los requisitos estándar, con controles adicionales y monitoreo continuo.'
        },
        '¿Se establecen y mantienen políticas para la clasificación y etiquetado de la información dentro del departamento de sistemas?': {
            1: 'No se establecen ni mantienen políticas para clasificación y etiquetado, lo que puede resultar en una gestión inadecuada de la información.',
            2: 'Políticas de clasificación y etiquetado establecidas pero no mantenidas adecuadamente, lo que puede llevar a inconsistencias.',
            3: 'Políticas de clasificación y etiquetado mantenidas regularmente, cumpliendo con la mayoría de los requisitos de ISO 27001.',
            4: 'Cumple totalmente con los requisitos de clasificación y etiquetado de ISO 27001, asegurando una gestión adecuada de la información.',
            5: 'Implementación avanzada que supera los requisitos estándar, con políticas detalladas y revisiones periódicas.'
        }
    },
    'Gestión de Incidentes de Seguridad de la Información': {
        '¿Existe un procedimiento documentado para la gestión de incidentes de seguridad de la información?': {
            1: 'No hay procedimiento documentado para la gestión de incidentes, lo que puede resultar en una respuesta inadecuada ante incidentes.',
            2: 'Procedimiento documentado pero no actualizado o implementado de manera limitada, lo que puede llevar a inconsistencias en la respuesta.',
            3: 'Procedimiento documentado y regularmente revisado e implementado, cumpliendo con la mayoría de los requisitos de ISO 27001.',
            4: 'Cumple totalmente con los requisitos de gestión de incidentes de ISO 27001, asegurando una respuesta adecuada y efectiva.',
            5: 'Implementación avanzada que supera los requisitos estándar, con procedimientos detallados y revisiones periódicas.'
        },
        '¿Se realiza capacitación y simulacros periódicos para el personal sobre cómo responder a incidentes de seguridad de la información?': {
            1: 'No se realizan capacitaciones ni simulacros sobre incidentes de seguridad, lo que puede resultar en una falta de preparación del personal.',
            2: 'Capacitaciones y simulacros realizados de manera irregular o insuficiente, lo que puede llevar a una respuesta inadecuada.',
            3: 'Capacitaciones y simulacros realizados regularmente, cumpliendo con la mayoría de los requisitos de ISO 27001.',
            4: 'Cumple totalmente con los requisitos de capacitación y simulacros de ISO 27001, asegurando una preparación adecuada del personal.',
            5: 'Implementación avanzada que supera los requisitos estándar, con capacitaciones y simulacros detallados y frecuentes.'
        }
    }
}


class EvaluacionISO27001(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Evaluación de Cumplimiento ISO 27001")
        self.geometry("800x600")
        self.calificaciones = {key: [] for key in rubricas.keys()}
        self.create_widgets()

    def create_widgets(self):
        self.tabs = ttk.Notebook(self)
        self.frames = {}

        for aspecto in rubricas.keys():
            frame = ttk.Frame(self.tabs)
            self.frames[aspecto] = frame
            self.tabs.add(frame, text=aspecto)
            self.create_aspect_widgets(frame, aspecto)

        self.tabs.pack(expand=1, fill='both')

        self.generate_button = tk.Button(self, text="Generar Informe", command=self.verificar_calificaciones)
        self.generate_button.pack(pady=20)

    def create_aspect_widgets(self, frame, aspecto):
        for pregunta in rubricas[aspecto].keys():
            label = tk.Label(frame, text=pregunta)
            label.pack(anchor='w', padx=10, pady=5)

            valores_descripciones = [desc for desc in rubricas[aspecto][pregunta].values()]
            combobox = ttk.Combobox(frame, values=valores_descripciones, width=100)
            combobox.pack(anchor='w', padx=10, pady=5)
            combobox.set("Seleccione una descripción")

            # Create a label to show the rating once selected
            rating_label = tk.Label(frame, text="Calificación: N/A")
            rating_label.pack(anchor='w', padx=10, pady=5)

            combobox.bind("<<ComboboxSelected>>",
                          lambda event, rb=rubricas[aspecto][pregunta], lbl=rating_label: self.update_rating(event, rb,
                                                                                                             lbl))
            self.calificaciones[aspecto].append((pregunta, combobox, rating_label))

    def update_rating(self, event, rubrica, label):
        selected_description = event.widget.get()
        for calificacion, descripcion in rubrica.items():
            if descripcion == selected_description:
                label.config(text=f"Calificación: {calificacion}")
                break

    def verificar_calificaciones(self):
        for aspecto, preguntas in self.calificaciones.items():
            for pregunta, combobox, rating_label in preguntas:
                if rating_label.cget("text") == "Calificación: N/A":
                    messagebox.showerror("Error de Calificación",
                                         f"No se ha evaluado todas las preguntas en la categoría '{aspecto}'.")
                    self.tabs.select(self.frames[aspecto])
                    return

        self.solicitar_datos_generar_informe()

    def solicitar_datos_generar_informe(self):
        self.nombre_compania = simpledialog.askstring("Nombre de la Compañía Evaluada",
                                                      "Ingrese el nombre de la compañía evaluada:")
        self.nombre_evaluador = simpledialog.askstring("Nombre y Apellido de la Persona que Evalúa",
                                                       "Ingrese el nombre y apellido de la persona que evalúa:")
        self.destinatario = simpledialog.askstring("Nombre y Apellido del Destinatario",
                                                   "Ingrese el nombre y apellido del destinatario del informe:")

        if not all([self.nombre_compania, self.nombre_evaluador, self.destinatario]):
            messagebox.showerror("Error", "Debe completar todos los campos para generar el informe.")
            return

        self.fecha_evaluacion = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        self.generar_informe()

    def generar_informe(self):
        calificaciones_input = {key: [] for key in rubricas.keys()}
        for aspecto, preguntas in self.calificaciones.items():
            for pregunta, combobox, rating_label in preguntas:
                try:
                    calificacion = int(rating_label.cget("text").split(": ")[1])
                    calificaciones_input[aspecto].append((pregunta, calificacion))
                except ValueError:
                    messagebox.showerror("Error de Calificación",
                                         "Por favor seleccione una descripción válida para todas las preguntas.")
                    return

        for aspecto, lista in calificaciones_input.items():
            if len(lista) != len(rubricas[aspecto]):
                messagebox.showerror("Error de Calificación",
                                     f"Debe calificar todas las preguntas en la categoría '{aspecto}' antes de generar el informe.")
                return

        promedios_ponderados, calificacion_final = procesar_calificaciones(calificaciones_input)
        generar_informe_word(calificaciones_input, promedios_ponderados, calificacion_final,
                             self.nombre_compania, self.nombre_evaluador,
                             self.destinatario, self.fecha_evaluacion)


def procesar_calificaciones(calificaciones):
    promedios = {aspecto: sum(valores[1] for valores in lista) / len(lista) for aspecto, lista in
                 calificaciones.items()}
    promedios_ponderados = {aspecto: (promedio / 5) * 20 for aspecto, promedio in promedios.items()}
    calificacion_final = sum(promedios_ponderados.values()) / len(promedios_ponderados) * 5
    return promedios_ponderados, calificacion_final


def generar_conclusion(calificacion_final):
    if 0 <= calificacion_final <= 25:
        return ("El departamento de sistemas muestra una falta significativa de cumplimiento en la gestión de acceso, "
                "seguridad física y ambiental, gestión de comunicaciones y operaciones, control de acceso a la información, "
                "y gestión de incidentes de seguridad de la información. No existen políticas ni procedimientos documentados, "
                "y los controles de seguridad son insuficientes o inexistentes, exponiendo la información a riesgos severos.")
    elif 26 <= calificacion_final <= 50:
        return (
            "El departamento de sistemas tiene algunos controles y políticas en su lugar, pero estos no son suficientemente robustos "
            "o no se aplican consistentemente. Existen políticas y procedimientos documentados en algunas áreas, pero pueden estar "
            "desactualizados o no ser efectivos en la práctica. Los controles de seguridad física y ambiental, así como las medidas "
            "de autenticación, se implementan de manera limitada, y las pruebas de vulnerabilidad y evaluaciones de riesgos se realizan "
            "de forma irregular.")
    elif 51 <= calificacion_final <= 75:
        return (
            "El departamento de sistemas ha implementado la mayoría de los controles de seguridad requeridos por la norma ISO 27001. "
            "Las políticas y procedimientos están documentados y se revisan regularmente. La seguridad física y ambiental es adecuada, "
            "y los controles de autenticación son robustos para la mayoría de los sistemas críticos. Las pruebas de vulnerabilidad y "
            "evaluaciones de riesgos se realizan de manera regular, aunque aún existen áreas que pueden mejorarse para alcanzar un nivel óptimo.")
    elif 76 <= calificacion_final <= 100:
        return (
            "El departamento de sistemas cumple completamente con los requisitos de la norma ISO 27001, y además implementa medidas adicionales "
            "que superan los estándares establecidos. Las políticas y procedimientos están completamente documentados y actualizados, y se revisan "
            "periódicamente. La seguridad física y ambiental es robusta y se monitorea continuamente. Los controles de autenticación incluyen medidas "
            "avanzadas como la autenticación multifactor, y las pruebas de vulnerabilidad y evaluaciones de riesgos se realizan de manera continua y "
            "exhaustiva. La gestión de incidentes es proactiva, con simulacros y capacitaciones regulares que aseguran una preparación adecuada del personal.")
    else:
        return "Calificación no válida."


def generar_informe_word(calificaciones, promedios_ponderados, calificacion_final, nombre_compania, nombre_evaluador,
                         destinatario, fecha_evaluacion):
    document = Document()

    # Carátula
    document.add_heading(
        'Informe de Autoevaluación de Cumplimiento de la Norma ISO 27001 (Sistema de Gestión de Seguridad de la Información)',
        0)
    document.add_paragraph(f'Compañía Evaluada: {nombre_compania}', style='Title')
    document.add_paragraph(f'Evaluador: {nombre_evaluador}', style='Heading 3')
    document.add_paragraph(f'Fecha de Evaluación: {fecha_evaluacion}', style='Heading 3')
    document.add_page_break()

    # Índice
    document.add_heading('Índice', level=1)
    document.add_paragraph('1. Carta de Introducción')
    document.add_paragraph('2. Resultados de la Autoevaluación')
    document.add_paragraph('3. Conclusión General')
    document.add_paragraph('4. Anexo I: Detalles de la Evaluación')
    document.add_paragraph('5. Gráficos')
    document.add_page_break()

    # Carta de introducción
    document.add_heading('Carta de Introducción', level=1)
    document.add_paragraph(f'{destinatario}', style='Heading 2')
    document.add_paragraph(f'{fecha_evaluacion.split()[0]}', style='Heading 3')
    document.add_paragraph()
    document.add_paragraph(f'Estimado/a {destinatario},')
    document.add_paragraph(
        f'Me dirijo a usted en mi calidad de usuario del sistema de autoevaluación de DeltechAudit para presentar un informe de mi autoevaluación del cumplimiento de la Norma ISO 27001 (Sistema de Gestión de Seguridad de la Información) para {nombre_compania}.'
    )
    document.add_heading('Limitación de Responsabilidad', level=2)
    document.add_paragraph(
        'DeltechAudit, como propietario del software utilizado para esta evaluación, se exime de cualquier responsabilidad derivada de la interpretación o uso de los resultados de este informe. Es importante destacar que este informe no ha sido auditado ni verificado por DeltechAudit y se proporciona únicamente como una herramienta orientativa para identificar posibles niveles de riesgo asociados con el cumplimiento de la Norma ISO 27001. Además, los controles evaluados en este informe representan solo una parte de los requeridos por la Norma ISO 27001.'
    )
    document.add_heading('Propósito del Informe', level=2)
    document.add_paragraph(
        'El objetivo principal de esta autoevaluación es identificar posibles áreas de riesgo y no conformidad dentro del sistema de gestión de seguridad de la información de su organización. La evaluación realizada permite obtener una visión preliminar de las fortalezas y debilidades en relación con los requisitos de la Norma ISO 27001.'
    )
    document.add_heading('Recomendación', level=2)
    document.add_paragraph(
        'Como resultado de esta evaluación preliminar, recomendamos encarecidamente llevar a cabo una auditoría exhaustiva por parte de profesionales calificados para valorar en detalle las vulnerabilidades e incumplimientos potenciales. Una auditoría completa permitirá implementar medidas correctivas adecuadas y asegurar que el sistema de gestión de seguridad de la información cumpla plenamente con la Norma ISO 27001.'
    )
    document.add_page_break()

    # Resultados de la autoevaluación
    document.add_heading('Resultados de la Autoevaluación', level=1)
    document.add_paragraph(
        'Los resultados detallados de la autoevaluación se incluyen en el Anexo I siguiente. Esperamos que los resultados proporcionados por esta herramienta sean útiles para que su organización comience a tomar muy en serio la seguridad de su información. Esta herramienta está diseñada para ayudar a identificar vulnerabilidades y a iniciar el proceso de minimizarlas.'
    )
    document.add_page_break()

    # Anexo I: Detalles de la evaluación
    document.add_heading('Anexo I: Detalles de la Evaluación', level=1)
    for aspecto, preguntas in calificaciones.items():
        document.add_heading(aspecto, level=2)
        for pregunta, calificacion in preguntas:
            descripcion = rubricas[aspecto][pregunta][calificacion]
            p = document.add_paragraph()
            p.add_run(f'{pregunta}: ').bold = True
            p.add_run(f'{calificacion} - {descripcion}')
        document.add_paragraph(f'Promedio del aspecto ({aspecto}): {promedios_ponderados[aspecto]:.2f} / 20')
        document.add_paragraph()
    document.add_paragraph(f'Calificación final del departamento de sistemas: {calificacion_final:.2f} / 100')
    document.add_page_break()

    # Conclusión general
    conclusion = generar_conclusion(calificacion_final)
    document.add_heading('Conclusión General', level=1)
    document.add_paragraph(conclusion)
    document.add_page_break()

    # Gráficos
    document.add_heading('Gráficos', level=1)

    # Gráfico de barras
    plt.figure(figsize=(10, 6))
    aspectos = list(promedios_ponderados.keys())
    valores = list(promedios_ponderados.values())
    plt.barh(aspectos, valores, color='skyblue')
    plt.xlabel('Nivel de Cumplimiento (sobre 20)')
    plt.title('Gráfico de Nivel de Cumplimiento por Dimensión')
    plt.savefig('grafico_barras.png')
    plt.close()

    # Insertar gráfico de barras
    document.add_paragraph('Gráfico de Nivel de Cumplimiento por Dimensión')
    document.add_picture('grafico_barras.png', width=Inches(6))
    document.add_page_break()

    # Gráfico de radar
    labels = list(promedios_ponderados.keys())
    stats = list(promedios_ponderados.values())

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    stats += stats[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, stats, color='skyblue', alpha=0.25)
    ax.plot(angles, stats, color='skyblue', linewidth=2)
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    plt.title('Gráfico de Radar por Dimensión')
    plt.savefig('grafico_radar.png')
    plt.close()

    # Insertar gráfico de radar
    document.add_paragraph('Gráfico de Radar por Dimensión')
    document.add_picture('grafico_radar.png', width=Inches(6))
    document.add_page_break()

    # Añadir texto al final del documento
    document.add_paragraph(
        'DeltechAudit es una Firma de Auditoría y Consultoría con 20 años de trayectoria. Podemos ayudarle a validar su autoauditoría. Escríbanos a ventas@smsauditores.ec')
    document.add_page_break()

    # Añadir pie de página
    for section in document.sections:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = f'www.deltechaudit.com - Fecha de Evaluación: {fecha_evaluacion}'
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Numerar páginas a partir de la segunda página
    for i, section in enumerate(document.sections):
        if i > 0:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = f'Página {i + 1}'
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save('Autoevaluacion_Norma_ISO_27001.docx')
    messagebox.showinfo("Informe Generado",
                        "El informe se ha generado correctamente en Autoevaluacion_Norma_ISO_27001.docx")


if __name__ == "__main__":
    app = EvaluacionISO27001()
    app.mainloop()
